#!/usr/bin/env python3
"""
每日思考 Telegram Bot

两种发布流程（缓冲区独立）：
  【每日记录】/d 切换模式 → 发文言文 → /daily 润色发布
  【学术思考】/t 切换模式 → 粘贴内容 → /publish 直接发布

指令：
  /d       — 切换到日常记录模式
  /t       — 切换到学术思考模式
  /daily   — 润色并发布日常记录
  /publish — 整理并发布学术思考
  /cancel  — 清空当前模式缓冲
"""

import os, json, base64, datetime, io, re
import urllib.request, urllib.error
from docx import Document
from openai import OpenAI
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler,
    CallbackQueryHandler, ContextTypes, filters
)

# ── 配置 ──────────────────────────────────────────────────────────
INSERT_MARK   = '<!-- ==================== ENTRIES ==================== -->'
THOUGHTS_PATH = 'thoughts.html'
DOCS_DIR      = 'files/thoughts'

def _load_config():
    cfg = {}
    path = os.path.expanduser('~/.thoughts_bot_config')
    if os.path.exists(path):
        for line in open(path):
            line = line.strip()
            if '=' in line and not line.startswith('#'):
                k, v = line.split('=', 1)
                cfg[k.strip()] = v.strip()
    for key in ('TELEGRAM_BOT_TOKEN', 'OPENAI_API_KEY', 'ALLOWED_USER_ID',
                'GITHUB_TOKEN', 'GITHUB_REPO'):
        if os.environ.get(key):
            cfg[key] = os.environ[key]
    return cfg

CFG        = _load_config()
BOT_TOKEN  = CFG.get('TELEGRAM_BOT_TOKEN', '')
OPENAI_KEY = CFG.get('OPENAI_API_KEY', '')
ALLOWED_ID = int(CFG.get('ALLOWED_USER_ID', '0'))
GH_TOKEN   = CFG.get('GITHUB_TOKEN', '')
GH_REPO    = CFG.get('GITHUB_REPO', '')

ai = OpenAI(api_key=OPENAI_KEY) if OPENAI_KEY else None

# ── 缓冲 ──────────────────────────────────────────────────────────
# sessions[uid] = {'daily': [...], 'thought': [...]}
sessions:      dict[int, dict[str, list[str]]] = {}
modes:         dict[int, str]  = {}   # 'daily' | 'thought'，默认 'daily'
drafts:        dict[int, tuple] = {}  # (mode, data)
title_editing: set[int]         = set()
# edit_state[uid] = {'entries': [...], 'html': str, 'sha': str, 'selected': int|None, 'field': str|None}
edit_state:    dict[int, dict]  = {}

def _sess(uid: int) -> dict[str, list[str]]:
    if uid not in sessions:
        sessions[uid] = {'daily': [], 'thought': []}
    return sessions[uid]

def _mode(uid: int) -> str:
    return modes.get(uid, 'daily')

def _mode_label(mode: str) -> str:
    return '日常记录' if mode == 'daily' else '学术思考'

# ── Prompts ───────────────────────────────────────────────────────
PROMPT_DAILY = """你是作者的文字助手。作者用文言文写了今天的日常记录，你需要：

1. 提炼一个简洁标题（文言风格，10字以内）
2. 将内容润色为简洁古雅的文言文，保留原意，控制在 80-120 字之间，1-2 段
3. 提取 2-4 个关键词作为标签

注意：全程使用简体中文，不得出现繁体字。

以 JSON 格式返回：
{"title": "...", "paragraphs": ["..."], "tags": ["..."]}
只返回 JSON。"""

PROMPT_THOUGHT = """你是作者的发布助手。作者已经写好了学术思考内容，你需要：

1. 提炼一个简洁标题（中文为主，可含英文，15字以内）
2. 写一段简要摘要（2-3句，概括核心论点，供网页展示）
3. 将正文原文照搬，按段落分开，供下载文档使用
4. 提取 3-5 个关键词作为标签（中英均可）

以 JSON 格式返回：
{"title": "...", "summary": "两三句摘要", "paragraphs": ["第一段原文", "第二段原文"], "tags": ["..."]}
只返回 JSON。"""

PROMPT_THOUGHT_FROM_DOC = """你是作者的发布助手。作者上传了一篇学术文档，你需要：

1. 提炼一个简洁标题（中文为主，可含英文，15字以内）
2. 写一段简要摘要（2-3句，概括核心论点，供网页展示）
3. 提取 3-5 个关键词作为标签（中英均可）

以 JSON 格式返回：
{"title": "...", "summary": "两三句摘要", "tags": ["..."]}
只返回 JSON。"""

def call_gpt(messages: list[str], prompt: str) -> dict:
    text = '\n\n'.join(messages)
    response = ai.chat.completions.create(
        model='gpt-4o',
        response_format={'type': 'json_object'},
        messages=[
            {'role': 'system', 'content': prompt},
            {'role': 'user',   'content': text},
        ],
    )
    return json.loads(response.choices[0].message.content)

# ── HTML ──────────────────────────────────────────────────────────
def build_article(data: dict, date_str: str, doc_url: str = None) -> str:
    if doc_url:
        body = f'<p>{data["summary"]}</p>\n      <p><a href="{doc_url}" download style="color:#888;font-size:13px;letter-spacing:1px;">↓ DOWNLOAD FULL TEXT</a></p>'
    else:
        body = '\n      '.join(f'<p>{p}</p>' for p in data['paragraphs'])
    tags  = '\n      '.join(f'<span class="tag">{t}</span>' for t in data.get('tags', []))
    tag_block = f'\n    <div class="entry-tags">\n      {tags}\n    </div>' if tags else ''
    return f'''
  <article class="entry" id="{date_str}">
    <div class="entry-header">
      <span class="entry-date">{date_str.replace("-", ".")}</span>
      <span class="entry-title">{data["title"]}</span>
    </div>
    <div class="entry-body">
      {body}
    </div>{tag_block}
  </article>
'''

# ── GitHub API ────────────────────────────────────────────────────
def _gh_request(method: str, path: str, body: dict = None):
    url  = f'https://api.github.com/repos/{GH_REPO}/contents/{path}'
    data = json.dumps(body).encode() if body else None
    req  = urllib.request.Request(url, data=data, method=method)
    req.add_header('Authorization', f'Bearer {GH_TOKEN}')
    req.add_header('Accept', 'application/vnd.github+json')
    req.add_header('Content-Type', 'application/json')
    with urllib.request.urlopen(req) as r:
        return json.loads(r.read())

def _parse_entries(html: str) -> list[dict]:
    """解析 HTML 中的所有条目，返回列表"""
    entries = []
    for m in re.finditer(r'<article class="entry" id="([^"]+)">(.*?)</article>', html, re.DOTALL):
        raw   = m.group(0)
        date  = m.group(1)
        inner = m.group(2)
        title_m = re.search(r'<span class="entry-title">([^<]+)</span>', inner)
        title = title_m.group(1) if title_m else ''
        # 正文文字（去掉 HTML 标签，仅供预览）
        body_m = re.search(r'<div class="entry-body">(.*?)</div>', inner, re.DOTALL)
        body_text = re.sub(r'<[^>]+>', '', body_m.group(1)).strip() if body_m else ''
        entries.append({'date': date, 'title': title, 'body_text': body_text, 'raw': raw})
    return entries

def _is_academic(entry: dict) -> bool:
    return bool(re.search(r'<a href="[^"]*" download', entry['raw']))

def _get_doc_path(entry: dict) -> str | None:
    """从条目 HTML 中提取 docx 在仓库里的路径"""
    m = re.search(r'href="https://yihuang\.art/([^"]+\.docx)"', entry['raw'])
    return m.group(1) if m else None

def _get_summary(entry: dict) -> str:
    """提取学术条目的摘要段（第一个非下载链接的 <p>）"""
    m = re.search(r'<p>(?!<a )(.*?)</p>', entry['raw'], re.DOTALL)
    return re.sub(r'<[^>]+>', '', m.group(1)).strip() if m else ''

def _update_entry_html(html: str, entry: dict, new_title: str = None, new_body: str = None, new_summary: str = None) -> str:
    """在 html 中就地替换指定条目的标题或正文"""
    old_raw = entry['raw']
    new_raw = old_raw
    if new_title:
        new_raw = new_raw.replace(
            f'<span class="entry-title">{entry["title"]}</span>',
            f'<span class="entry-title">{new_title}</span>'
        )
    if new_body is not None:
        # 保留下载链接（学术条目）
        dl_m = re.search(r'<p><a href="[^"]*" download[^>]*>.*?</a></p>', old_raw, re.DOTALL)
        dl_link = f'\n      {dl_m.group(0)}' if dl_m else ''
        new_paras = '\n      '.join(f'<p>{p.strip()}</p>' for p in new_body.split('\n') if p.strip())
        new_raw = re.sub(
            r'(<div class="entry-body">).*?(</div>)',
            lambda m: m.group(1) + f'\n      {new_paras}{dl_link}\n    ' + m.group(2),
            new_raw, flags=re.DOTALL
        )
    if new_summary is not None:
        new_raw = re.sub(
            r'<p>(?!<a )(.*?)</p>',
            f'<p>{new_summary}</p>',
            new_raw, count=1, flags=re.DOTALL
        )
    return html.replace(old_raw, new_raw)

def github_update(article: str, date_str: str) -> tuple[bool, str]:
    try:
        resp     = _gh_request('GET', THOUGHTS_PATH)
        sha      = resp['sha']
        html     = base64.b64decode(resp['content']).decode('utf-8')
        if INSERT_MARK not in html:
            return False, '找不到 thoughts.html 里的标记位置'
        new_html = html.replace(INSERT_MARK, INSERT_MARK + '\n' + article)
        _gh_request('PUT', THOUGHTS_PATH, {
            'message': f'thoughts: {date_str}',
            'content': base64.b64encode(new_html.encode('utf-8')).decode(),
            'sha': sha,
        })
        return True, ''
    except urllib.error.HTTPError as e:
        return False, f'GitHub API 错误 {e.code}: {e.read().decode()}'
    except Exception as e:
        return False, str(e)

def _set_font(run, size_pt):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from lxml import etree
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), 'SimSun')

def github_upload_docx(data: dict, date_str: str, docx_bytes: bytes = None) -> tuple[bool, str]:
    try:
        if docx_bytes is None:
            from docx.shared import Pt
            doc = Document()
            title_para = doc.add_paragraph()
            run = title_para.add_run(data['title'])
            run.bold = True
            _set_font(run, 16)
            for para in data.get('paragraphs', []):
                p = doc.add_paragraph()
                run = p.add_run(para)
                _set_font(run, 12)
            buf = io.BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()

        ts   = datetime.datetime.now().strftime('%H%M%S')
        path = f'{DOCS_DIR}/{date_str}-{ts}.docx'
        try:
            existing = _gh_request('GET', path)
            sha = existing['sha']
        except urllib.error.HTTPError:
            sha = None

        body = {
            'message': f'thoughts-doc: {date_str}',
            'content': base64.b64encode(docx_bytes).decode(),
        }
        if sha:
            body['sha'] = sha
        _gh_request('PUT', path, body)
        url = f'https://yihuang.art/{path}'
        return True, url
    except urllib.error.HTTPError as e:
        return False, f'GitHub API 错误 {e.code}: {e.read().decode()}'
    except Exception as e:
        return False, str(e)

# ── Keyboards ─────────────────────────────────────────────────────
def _date_picker_keyboard(code: str) -> InlineKeyboardMarkup:
    """code: 'd' for daily, 't' for thought"""
    today = datetime.date.today()
    labels = ['今天', '昨天', '前天']
    row = [
        InlineKeyboardButton(
            f'{labels[i]} {(today - datetime.timedelta(days=i)).strftime("%m.%d")}',
            callback_data=f'dpick_{code}_{i}'
        )
        for i in range(3)
    ]
    return InlineKeyboardMarkup([row, [InlineKeyboardButton('❌ 取消', callback_data='cancel')]])

def _daily_keyboard():
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton('✅ 发布',     callback_data='daily_confirm'),
            InlineKeyboardButton('🔄 重新生成', callback_data='daily_retry'),
        ],
        [
            InlineKeyboardButton('✏️ 改标题',   callback_data='edit_title'),
            InlineKeyboardButton('❌ 清空',     callback_data='cancel'),
        ],
    ])

def _thought_keyboard():
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton('✅ 发布',   callback_data='thought_confirm'),
            InlineKeyboardButton('✏️ 改标题', callback_data='edit_title'),
        ],
        [
            InlineKeyboardButton('❌ 清空',   callback_data='cancel'),
        ],
    ])

# ── 预览辅助 ──────────────────────────────────────────────────────
def _preview_text(data: dict) -> str:
    """每日记录预览：显示完整正文（较短）"""
    tags_str = '  '.join(f'#{t}' for t in data.get('tags', []))
    text = f'📝 《{data["title"]}》\n\n' + '\n\n'.join(data['paragraphs'])
    if tags_str:
        text += f'\n\n{tags_str}'
    return text[:4000]  # Telegram 上限保护

def _thought_preview_text(data: dict) -> str:
    """学术思考预览：只显示标题 + 摘要 + 标签，全文进 docx 不展示"""
    tags_str = '  '.join(f'#{t}' for t in data.get('tags', []))
    text = f'📝 《{data["title"]}》\n\n{data.get("summary", "")}'
    if tags_str:
        text += f'\n\n{tags_str}'
    text += '\n\n（全文已整理，发布后可下载完整文档）'
    return text

# ── Handlers ──────────────────────────────────────────────────────
def is_allowed(update: Update) -> bool:
    return ALLOWED_ID == 0 or update.effective_user.id == ALLOWED_ID

async def cmd_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update): return
    uid = update.effective_user.id
    _sess(uid)
    await update.message.reply_text(
        '你好！\n\n'
        '两个缓冲区独立，互不干扰：\n\n'
        '/d  — 切换到【日常记录】模式\n'
        '/t  — 切换到【学术思考】模式\n\n'
        '切换后发送内容，再用：\n'
        '/daily   — 润色发布日常记录\n'
        '/publish — 整理发布学术思考\n'
        '/edit    — 查看并修改最近条目\n'
        '/cancel  — 清空当前模式缓冲'
    )

async def cmd_d(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """切换到日常记录模式"""
    if not is_allowed(update): return
    uid = update.effective_user.id
    modes[uid] = 'daily'
    count = len(_sess(uid)['daily'])
    await update.message.reply_text(
        f'📋 已切换到【日常记录】模式\n'
        f'当前缓冲：{count} 条\n\n'
        f'发送文言文内容，完成后 /daily 润色发布。'
    )

async def cmd_t(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """切换到学术思考模式"""
    if not is_allowed(update): return
    uid = update.effective_user.id
    modes[uid] = 'thought'
    count = len(_sess(uid)['thought'])
    await update.message.reply_text(
        f'📚 已切换到【学术思考】模式\n'
        f'当前缓冲：{count} 条\n\n'
        f'粘贴内容，完成后 /publish 整理发布。'
    )

async def cmd_cancel(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update): return
    uid = update.effective_user.id
    mode = _mode(uid)
    _sess(uid)[mode] = []
    drafts.pop(uid, None)
    title_editing.discard(uid)
    await update.message.reply_text(f'已清空【{_mode_label(mode)}】缓冲，重新开始吧。')

async def cmd_edit(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """/edit — 查看最近一天的条目并修改"""
    if not is_allowed(update): return
    uid = update.effective_user.id
    await update.message.reply_text('读取中…')
    try:
        resp    = _gh_request('GET', THOUGHTS_PATH)
        sha     = resp['sha']
        html    = base64.b64decode(resp['content']).decode('utf-8')
        entries = _parse_entries(html)
        if not entries:
            await update.message.reply_text('没有找到任何条目。')
            return
        # 取三日内的条目
        cutoff = (datetime.date.today() - datetime.timedelta(days=2)).strftime('%Y-%m-%d')
        recent = [e for e in entries if e['date'] >= cutoff]
        if not recent:
            recent = [entries[0]]  # 至少显示最新一条
        edit_state[uid] = {'entries': recent, 'html': html, 'sha': sha, 'selected': None, 'field': None}
        text = '📅 三日内的条目，选择要编辑的：'
        buttons = []
        # 缓冲区草稿入口
        sess = _sess(uid)
        if sess['thought']:
            buttons.append([InlineKeyboardButton(f'📚 学术草稿（{len(sess["thought"])}条）→ 整理发布', callback_data='draft_thought')])
        if sess['daily']:
            buttons.append([InlineKeyboardButton(f'📋 日常草稿（{len(sess["daily"])}条）→ 润色发布', callback_data='draft_daily')])
        buttons += [[InlineKeyboardButton(f'{e["date"].replace("-",".")}  《{e["title"]}》', callback_data=f'esel_{i}')]
                    for i, e in enumerate(recent)]
        buttons.append([InlineKeyboardButton('❌ 取消', callback_data='edit_cancel')])
        await update.message.reply_text(text, reply_markup=InlineKeyboardMarkup(buttons))
    except Exception as e:
        await update.message.reply_text(f'❌ 出错了：{e}')

async def cmd_daily(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update): return
    uid  = update.effective_user.id
    msgs = _sess(uid)['daily']
    if not msgs:
        await update.message.reply_text('日常记录缓冲区为空。\n请先 /d 切换到日常模式，再发送内容。')
        return
    await update.message.reply_text('润色中…')
    try:
        data = call_gpt(msgs, PROMPT_DAILY)
        drafts[uid] = ('daily', data)
        await update.message.reply_text(_preview_text(data), reply_markup=_daily_keyboard())
    except Exception as e:
        await update.message.reply_text(f'❌ 出错了：{e}')

async def cmd_publish(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update): return
    uid  = update.effective_user.id
    msgs = _sess(uid)['thought']
    if not msgs:
        await update.message.reply_text('学术思考缓冲区为空。\n请先 /t 切换到学术模式，再粘贴内容。')
        return
    await update.message.reply_text('整理中…')
    try:
        data = call_gpt(msgs, PROMPT_THOUGHT)
        drafts[uid] = ('thought', data)
        await update.message.reply_text(_thought_preview_text(data), reply_markup=_thought_keyboard())
    except Exception as e:
        await update.message.reply_text(f'❌ 出错了：{e}')

async def on_document(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """接收 .docx 文件 → 提取文字 → GPT 拟标题摘要 → 预览确认"""
    if not is_allowed(update): return
    uid = update.effective_user.id
    doc_msg = update.message.document
    if not doc_msg:
        return
    fname = doc_msg.file_name or ''
    if not fname.lower().endswith('.docx'):
        await update.message.reply_text('请发送 .docx 格式的 Word 文档。')
        return
    await update.message.reply_text('读取文档中…')
    try:
        tg_file = await ctx.bot.get_file(doc_msg.file_id)
        buf = io.BytesIO()
        await tg_file.download_to_memory(buf)
        docx_bytes = buf.getvalue()
        buf.seek(0)
        word_doc = Document(buf)
        full_text = '\n'.join(p.text for p in word_doc.paragraphs if p.text.strip())
        if not full_text:
            await update.message.reply_text('文档内容为空，无法处理。')
            return
        data = call_gpt([full_text], PROMPT_THOUGHT_FROM_DOC)
        drafts[uid] = ('thought_doc', data, docx_bytes)
        modes[uid] = 'thought'
        await update.message.reply_text(_thought_preview_text(data), reply_markup=_thought_keyboard())
    except Exception as e:
        await update.message.reply_text(f'❌ 出错了：{e}')

async def on_button(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    uid = query.from_user.id
    if ALLOWED_ID != 0 and uid != ALLOWED_ID:
        return

    action = query.data

    if action in ('draft_thought', 'draft_daily'):
        mode = 'thought' if action == 'draft_thought' else 'daily'
        msgs = _sess(uid)[mode]
        if not msgs:
            await query.edit_message_text('缓冲区已空，请先发送内容。')
            return
        await query.edit_message_text('整理中…')
        try:
            prompt = PROMPT_THOUGHT if mode == 'thought' else PROMPT_DAILY
            data   = call_gpt(msgs, prompt)
            drafts[uid] = (mode, data)
            if mode == 'thought':
                await ctx.bot.send_message(uid, _thought_preview_text(data), reply_markup=_thought_keyboard())
            else:
                await ctx.bot.send_message(uid, _preview_text(data), reply_markup=_daily_keyboard())
        except Exception as e:
            await ctx.bot.send_message(uid, f'❌ 出错了：{e}')

    elif action == 'edit_cancel':
        edit_state.pop(uid, None)
        await query.edit_message_text('已取消编辑。')

    elif action.startswith('esel_'):
        state = edit_state.get(uid)
        if not state:
            await query.edit_message_text('会话已过期，请重发 /edit。')
            return
        idx = int(action[5:])
        state['selected'] = idx
        entry = state['entries'][idx]
        preview = entry['body_text'][:200] + ('…' if len(entry['body_text']) > 200 else '')
        text = f'《{entry["title"]}》\n\n{preview}\n\n选择要修改的内容：'
        if _is_academic(entry):
            buttons = [
                [InlineKeyboardButton('✏️ 改标题', callback_data='efield_title'),
                 InlineKeyboardButton('✏️ 改摘要', callback_data='efield_summary')],
                [InlineKeyboardButton('✏️ 改全文', callback_data='efield_docx')],
                [InlineKeyboardButton('🗑️ 删除此条目', callback_data='edelete')],
                [InlineKeyboardButton('← 返回',   callback_data='edit_back'),
                 InlineKeyboardButton('❌ 取消',   callback_data='edit_cancel')],
            ]
        else:
            buttons = [
                [InlineKeyboardButton('✏️ 改标题', callback_data='efield_title'),
                 InlineKeyboardButton('✏️ 改正文', callback_data='efield_body')],
                [InlineKeyboardButton('🗑️ 删除此条目', callback_data='edelete')],
                [InlineKeyboardButton('← 返回',   callback_data='edit_back'),
                 InlineKeyboardButton('❌ 取消',   callback_data='edit_cancel')],
            ]
        await query.edit_message_text(text, reply_markup=InlineKeyboardMarkup(buttons))

    elif action == 'edelete':
        state = edit_state.get(uid)
        if not state or state['selected'] is None:
            await query.edit_message_text('会话已过期，请重发 /edit。')
            return
        entry = state['entries'][state['selected']]
        await query.edit_message_text(
            f'⚠️ 确认删除《{entry["title"]}》（{entry["date"].replace("-",".")}）？\n\n此操作不可撤销。',
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton('🗑️ 确认删除', callback_data='edelete_confirm'),
                 InlineKeyboardButton('← 取消',     callback_data=f'esel_{state["selected"]}')],
            ])
        )

    elif action == 'edelete_confirm':
        state = edit_state.get(uid)
        if not state or state['selected'] is None:
            await query.edit_message_text('会话已过期，请重发 /edit。')
            return
        entry = state['entries'][state['selected']]
        html  = state['html']
        sha   = state['sha']
        await query.edit_message_text('删除中…')
        try:
            new_html = html.replace(entry['raw'], '')
            _gh_request('PUT', THOUGHTS_PATH, {
                'message': f'thoughts-delete: {entry["date"]}',
                'content': base64.b64encode(new_html.encode('utf-8')).decode(),
                'sha': sha,
            })
            # 如果是学术条目，一并删除 docx
            doc_path = _get_doc_path(entry)
            if doc_path:
                try:
                    doc_info = _gh_request('GET', doc_path)
                    _gh_request('DELETE', doc_path, {
                        'message': f'thoughts-doc-delete: {entry["date"]}',
                        'sha': doc_info['sha'],
                    })
                except Exception:
                    pass  # docx 删除失败不影响主流程
            edit_state.pop(uid, None)
            await query.edit_message_text(f'✅ 已删除《{entry["title"]}》')
        except Exception as e:
            await query.edit_message_text(f'❌ 删除失败：{e}')

    elif action == 'edit_back':
        state = edit_state.get(uid)
        if not state:
            await query.edit_message_text('会话已过期，请重发 /edit。')
            return
        state['selected'] = None
        state['field'] = None
        entries = state['entries']
        date_label = entries[0]['date'].replace('-', '.')
        text = f'📅 {date_label} 的条目，选择要编辑的：'
        buttons = [[InlineKeyboardButton(f'《{e["title"]}》', callback_data=f'esel_{i}')]
                   for i, e in enumerate(entries)]
        buttons.append([InlineKeyboardButton('❌ 取消', callback_data='edit_cancel')])
        await query.edit_message_text(text, reply_markup=InlineKeyboardMarkup(buttons))

    elif action in ('efield_title', 'efield_body', 'efield_summary', 'efield_docx'):
        state = edit_state.get(uid)
        if not state or state['selected'] is None:
            await query.edit_message_text('会话已过期，请重发 /edit。')
            return
        field = {'efield_title': 'title', 'efield_body': 'body', 'efield_summary': 'summary', 'efield_docx': 'docx'}[action]
        state['field'] = field
        entry = state['entries'][state['selected']]
        if field == 'title':
            await query.edit_message_text(f'当前标题：《{entry["title"]}》\n\n请发送新标题：')
        elif field == 'summary':
            cur = _get_summary(entry)
            await query.edit_message_text(f'当前摘要：\n{cur}\n\n请发送新摘要：')
        elif field == 'docx':
            await query.edit_message_text(
                f'请发送新的全文内容（多段用换行分隔），将重新生成 Word 文档并上传：'
            )
        else:
            await query.edit_message_text(
                f'当前正文：\n{entry["body_text"][:300]}{"…" if len(entry["body_text"])>300 else ""}\n\n'
                f'请发送新正文（多段用换行分隔）：'
            )

    elif action == 'cancel':
        draft = drafts.get(uid)
        if draft:
            mode, _ = draft
            _sess(uid)[mode] = []
        drafts.pop(uid, None)
        title_editing.discard(uid)
        await query.edit_message_text('已清空，重新开始吧。')

    elif action == 'edit_title':
        draft = drafts.get(uid)
        if not draft:
            await query.edit_message_text('草稿已过期。')
            return
        _, data = draft
        title_editing.add(uid)
        await query.edit_message_text(
            f'当前标题：《{data["title"]}》\n\n'
            f'请发送新标题：'
        )

    elif action == 'daily_confirm':
        draft = drafts.get(uid)
        if not draft or not isinstance(draft, tuple):
            await query.edit_message_text('草稿已过期，请重发 /daily。')
            return
        await query.edit_message_text('选择发布日期：', reply_markup=_date_picker_keyboard('d'))

    elif action == 'daily_retry':
        msgs = _sess(uid)['daily']
        if not msgs:
            await query.edit_message_text('消息已丢失，请重新发送内容。')
            return
        await query.edit_message_text('重新润色中…')
        try:
            data = call_gpt(msgs, PROMPT_DAILY)
            drafts[uid] = ('daily', data)
            await ctx.bot.send_message(uid, _preview_text(data), reply_markup=_daily_keyboard())
        except Exception as e:
            await ctx.bot.send_message(uid, f'❌ 出错了：{e}')

    elif action == 'thought_confirm':
        draft = drafts.get(uid)
        if not draft or not isinstance(draft, tuple):
            await query.edit_message_text('草稿已过期，请重发 /publish。')
            return
        await query.edit_message_text('选择发布日期：', reply_markup=_date_picker_keyboard('t'))

    elif action.startswith('dpick_'):
        # dpick_d_0 / dpick_t_1 etc.
        _, code, offset_s = action.split('_')
        offset   = int(offset_s)
        date_str = (datetime.date.today() - datetime.timedelta(days=offset)).strftime('%Y-%m-%d')
        draft    = drafts.get(uid)
        if not draft:
            await query.edit_message_text('草稿已过期，请重新生成。')
            return
        if code == 'd':
            _, data = draft
            await query.edit_message_text('发布中…')
            ok, err = github_update(build_article(data, date_str), date_str)
            if ok:
                _sess(uid)['daily'] = []
                drafts.pop(uid, None)
                await query.edit_message_text(
                    f'✅ 已发布《{data["title"]}》\nhttps://yihuang.art/thoughts.html#{date_str}')
            else:
                await query.edit_message_text(f'❌ 发布失败：{err}')
        else:
            if draft[0] == 'thought_doc':
                _, data, docx_bytes = draft
            else:
                _, data = draft
                docx_bytes = None
            await query.edit_message_text('上传文档中…')
            ok, doc_url = github_upload_docx(data, date_str, docx_bytes)
            if not ok:
                await query.edit_message_text(f'❌ 文档上传失败：{doc_url}')
                return
            ok, err = github_update(build_article(data, date_str, doc_url), date_str)
            if ok:
                _sess(uid)['thought'] = []
                drafts.pop(uid, None)
                await query.edit_message_text(
                    f'✅ 已发布《{data["title"]}》\nhttps://yihuang.art/thoughts.html#{date_str}')
            else:
                await query.edit_message_text(f'❌ 发布失败：{err}')

async def on_message(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update): return
    uid  = update.effective_user.id
    text = update.message.text.strip()
    if not text:
        return

    # /edit 编辑状态
    state = edit_state.get(uid)
    if state and state.get('field') and state.get('selected') is not None:
        field   = state['field']
        entry   = state['entries'][state['selected']]
        html    = state['html']
        sha     = state['sha']
        if field == 'docx':
            # 重新生成 Word 文档并上传
            doc_path = _get_doc_path(entry)
            if not doc_path:
                await update.message.reply_text('❌ 找不到文档路径。')
                state['field'] = None
                return
            try:
                doc = Document()
                title_para = doc.add_paragraph()
                run = title_para.add_run(entry['title'])
                run.bold = True
                _set_font(run, 16)
                for para_text in text.split('\n'):
                    if para_text.strip():
                        p = doc.add_paragraph()
                        run = p.add_run(para_text.strip())
                        _set_font(run, 12)
                buf = io.BytesIO()
                doc.save(buf)
                docx_bytes = buf.getvalue()
                try:
                    existing = _gh_request('GET', doc_path)
                    doc_sha = existing['sha']
                except urllib.error.HTTPError:
                    doc_sha = None
                doc_body = {'message': f'thoughts-doc-edit: {entry["date"]}',
                            'content': base64.b64encode(docx_bytes).decode()}
                if doc_sha:
                    doc_body['sha'] = doc_sha
                _gh_request('PUT', doc_path, doc_body)
                state['field'] = None
                await update.message.reply_text(f'✅ 全文文档已更新！\nhttps://yihuang.art/thoughts.html#{entry["date"]}')
            except Exception as e:
                await update.message.reply_text(f'❌ 更新失败：{e}')
            return

        new_title   = text if field == 'title'   else None
        new_body    = text if field == 'body'    else None
        new_summary = text if field == 'summary' else None
        new_html    = _update_entry_html(html, entry, new_title=new_title, new_body=new_body, new_summary=new_summary)
        try:
            _gh_request('PUT', THOUGHTS_PATH, {
                'message': f'thoughts-edit: {entry["date"]}',
                'content': base64.b64encode(new_html.encode('utf-8')).decode(),
                'sha': sha,
            })
            if field == 'title':
                entry['title'] = text
            else:
                entry['body_text'] = text
            state['field'] = None
            state['html']  = new_html
            label = '标题' if field == 'title' else '正文'
            await update.message.reply_text(f'✅ {label}已更新！\nhttps://yihuang.art/thoughts.html#{entry["date"]}')
        except Exception as e:
            await update.message.reply_text(f'❌ 更新失败：{e}')
        return

    # 标题编辑状态：下一条消息作为新标题
    if uid in title_editing:
        title_editing.discard(uid)
        draft = drafts.get(uid)
        if draft:
            mode, data = draft
            data['title'] = text
            drafts[uid] = (mode, data)
            keyboard = _daily_keyboard() if mode == 'daily' else _thought_keyboard()
            preview  = _preview_text(data) if mode == 'daily' else _thought_preview_text(data)
            await update.message.reply_text(preview, reply_markup=keyboard)
        return

    # 正常消息：路由到当前模式的缓冲区
    mode = _mode(uid)
    _sess(uid)[mode].append(text)
    await update.message.reply_text(
        f'✓ 收到（{_mode_label(mode)}）\n'
        f'/d 日常模式  /t 学术模式  /daily 润色  /publish 发布'
    )

# ── 主入口 ────────────────────────────────────────────────────────
def main():
    if not BOT_TOKEN:
        print('❌ 缺少 TELEGRAM_BOT_TOKEN'); return
    if not OPENAI_KEY:
        print('❌ 缺少 OPENAI_API_KEY'); return
    if not GH_TOKEN or not GH_REPO:
        print('❌ 缺少 GITHUB_TOKEN 或 GITHUB_REPO'); return
    if not ALLOWED_ID:
        print('⚠️  ALLOWED_USER_ID 未设置，任何人都能使用此 bot')

    app = ApplicationBuilder().token(BOT_TOKEN).build()
    app.add_handler(CommandHandler('start',   cmd_start))
    app.add_handler(CommandHandler('d',       cmd_d))
    app.add_handler(CommandHandler('t',       cmd_t))
    app.add_handler(CommandHandler('cancel',  cmd_cancel))
    app.add_handler(CommandHandler('daily',   cmd_daily))
    app.add_handler(CommandHandler('publish', cmd_publish))
    app.add_handler(CommandHandler('edit',    cmd_edit))
    app.add_handler(CallbackQueryHandler(on_button))
    app.add_handler(MessageHandler(filters.Document.ALL, on_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, on_message))

    print('Bot 运行中…  Ctrl+C 停止')
    app.run_polling()

if __name__ == '__main__':
    main()
